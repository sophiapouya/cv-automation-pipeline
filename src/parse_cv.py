#!/usr/bin/env python3
"""
Stage 1 – Parse a BCM-style academic CV (PDF or .docx/.doc) and produce
a baseline publications CSV.

Usage:
    python parse_cv.py --cv path/to/cv.pdf --out publications.csv
    python parse_cv.py --cv path/to/cv.docx --out publications.csv

Outputs a CSV with columns:
    id, title, authors, year, journal, doi, pmid, openalex_id,
    status, pub_type, source, notes
"""

import argparse
import csv
import hashlib
import re
import subprocess
import sys
import tempfile
from pathlib import Path

# ── CSV schema ────────────────────────────────────────────────────────────────
CSV_FIELDS = [
    "id",
    "title",
    "authors",
    "year",
    "journal",
    "doi",
    "pmid",
    "openalex_id",
    "status",
    "pub_type",
    "source",
    "notes",
]

DOI_RE = re.compile(r"10\.\d{4,9}/\S+", re.I)
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
PMID_RE = re.compile(r"PMID[:\s]+(\d+)", re.I)
WS_RE = re.compile(r"\s+")

SECTION_MARKERS = [
    ("peer_reviewed_published", re.compile(r"^\s*a\.\s*Published\s*$", re.I)),
    ("peer_reviewed_in_press", re.compile(r"^\s*b\.\s*Accepted\s*/?\s*In Press\s*$", re.I)),
    ("no_peer_review", re.compile(r"^\s*2\.\s*Full Papers Without Peer Review\b", re.I)),
    ("in_preparation", re.compile(r"^\s*b\.\s*In Preparation\s*$", re.I)),
    ("book_chapters", re.compile(r"^\s*c\.\s*Book Chapters Written\b", re.I)),
    ("patents", re.compile(r"^\s*1\.\s*Patents\b", re.I)),
]

SECTION_CONFIG = {
    "peer_reviewed_published": ("published", "peer_reviewed"),
    "peer_reviewed_in_press": ("in_press", "peer_reviewed"),
    "book_chapters": ("published", "book_chapter"),
    "in_preparation": ("preprint", "peer_reviewed"),
    "no_peer_review": ("published", "non_peer_reviewed"),
}

SKIP_LINE_PATTERNS = [
    re.compile(r"^\s*(C\.\s*Publications|4\.\s*Books|5\.\s*Other Works)\b", re.I),
    re.compile(r"^\s*[abc]\.\s*(Published|Accepted\s*/?\s*In Press|Complete Books Written|Books Edited)\b", re.I),
    re.compile(r"^\s*D\.\s*Innovation and Commercialization\b", re.I),
    re.compile(r"^\s*None\s*$", re.I),
    re.compile(r"^\s*NOTE:\b", re.I),
    re.compile(r"^\s*Other Works Communicating Research Results\b", re.I),
]

SURNAME_FIRST_TOKEN = re.compile(
    r"""
    (?P<author>
        [^\W\d_][\w'’.\-]+(?:\s+[^\W\d_][\w'’.\-]+)*,\s*
        [A-Z][A-Za-z.\-*]*
    )
    """,
    re.X | re.UNICODE,
)


def make_id(title: str) -> str:
    norm = re.sub(r"\W+", " ", title.lower()).strip()
    return hashlib.sha1(norm.encode()).hexdigest()[:8]


def clean(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u2028", " ").replace("\u2029", " ")
    text = text.replace("​", "").replace("‌", "").replace("‍", "").replace("⁠", "").replace("­", "")
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    return WS_RE.sub(" ", text).strip(" \t.;")


def read_with_textutil(cv_path: Path, fmt: str = "txt") -> str:
    result = subprocess.run(
        ["textutil", "-convert", fmt, "-stdout", str(cv_path)],
        capture_output=True,
        text=True,
        timeout=60,
        check=False,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "textutil conversion failed")
    return result.stdout


def get_text(cv_path: Path) -> str:
    suffix = cv_path.suffix.lower()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(cv_path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            sys.exit(f"[ERROR] PDF read failed: {exc}")

    if suffix == ".docx":
        try:
            from docx import Document

            doc = Document(str(cv_path))
            paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
            return "\n".join(paragraphs)
        except Exception as exc:
            sys.exit(f"[ERROR] DOCX read failed: {exc}")

    if suffix == ".doc":
        tmpdir = tempfile.mkdtemp()
        lo_cmd = "soffice" if sys.platform == "win32" else "libreoffice"
        result = subprocess.run(
            [lo_cmd, "--headless", "--convert-to", "docx", "--outdir", tmpdir, str(cv_path)],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        if result.returncode != 0:
            sys.exit(f"[ERROR] LibreOffice conversion failed:\n{result.stderr}")
        work_path = Path(tmpdir) / f"{cv_path.stem}.docx"
        try:
            from docx import Document

            doc = Document(str(work_path))
            paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
            return "\n".join(paragraphs)
        except Exception as exc:
            sys.exit(f"[ERROR] DOCX read failed: {exc}")

    return cv_path.read_text(encoding="utf-8", errors="ignore")


def normalize_lines(text: str) -> list[str]:
    return [clean(line) for line in text.splitlines()]


def split_into_sections(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {name: [] for name, _ in SECTION_MARKERS}
    current = None
    for raw_line in normalize_lines(text):
        if not raw_line:
            if current:
                sections[current].append("")
            continue

        matched = next((name for name, pat in SECTION_MARKERS if pat.search(raw_line)), None)
        if matched:
            current = matched
            continue

        if current:
            sections[current].append(raw_line)
    return sections


def should_skip_line(line: str) -> bool:
    return any(pat.search(line) for pat in SKIP_LINE_PATTERNS)


def looks_like_entry_start(line: str) -> bool:
    if not line or should_skip_line(line):
        return False
    if DOI_RE.search(line) or "Accepted for publication" in line or "In review" in line:
        return True
    if YEAR_RE.search(line) and "," in line:
        return True
    return False


def collect_entry_blocks(lines: list[str]) -> list[str]:
    blocks: list[str] = []
    current = ""

    for line in lines:
        if should_skip_line(line):
            continue
        if not line:
            if current:
                blocks.append(clean(current))
                current = ""
            continue
        if current and looks_like_entry_start(line):
            blocks.append(clean(current))
            current = line
            continue
        current = f"{current} {line}".strip() if current else line

    if current:
        blocks.append(clean(current))
    return [block for block in blocks if block]


def parse_doi(block: str) -> str:
    match = DOI_RE.search(block)
    return match.group(0).rstrip(".,;)") if match else ""


def parse_year(block: str) -> str:
    match = re.search(r"\((19|20)\d{2}\)", block)
    if match:
        return match.group(0).strip("()")
    match = YEAR_RE.search(block)
    return match.group(0) if match else ""


def split_surname_first_authors(text: str) -> tuple[str, str]:
    pos = 0
    matches = []
    while True:
        match = SURNAME_FIRST_TOKEN.match(text, pos)
        if not match:
            break
        matches.append(match.group("author"))
        pos = match.end()
        sep = re.match(r"(?:,\s*| and\s+)", text[pos:])
        if not sep:
            break
        pos += sep.end()

    if matches and pos < len(text):
        return clean(text[:pos]), clean(text[pos:])
    return "", text


def looks_like_person_name(chunk: str) -> bool:
    chunk = clean(chunk).rstrip(".")
    if not chunk or any(word.islower() for word in chunk.split()):
        return False
    words = chunk.split()
    if len(words) < 2 or len(words) > 4:
        return False
    for word in words:
        if re.fullmatch(r"[A-Z]\.?", word):
            continue
        if not re.fullmatch(r"[A-Z][A-Za-z'’.\-]+", word):
            return False
    return True


def split_full_name_authors(text: str) -> tuple[str, str]:
    parts = [part.strip() for part in text.split(",")]
    authors = []
    consumed = 0
    for part in parts:
        if not looks_like_person_name(part):
            break
        authors.append(part)
        consumed += len(part)
        if consumed < len(text) and text[consumed:consumed + 1] == ",":
            consumed += 1
        while consumed < len(text) and text[consumed] == " ":
            consumed += 1
    if authors and consumed < len(text):
        return clean(", ".join(authors)), clean(text[consumed:])
    return "", text


def split_authors_and_rest(text: str) -> tuple[str, str]:
    authors, rest = split_surname_first_authors(text)
    if authors:
        return authors, rest
    authors, rest = split_full_name_authors(text)
    if authors:
        return authors, rest
    return "", text


def parse_remainder(rest: str, status: str) -> tuple[str, str]:
    rest = clean(rest)
    if not rest:
        return "", ""

    for marker, label in (
        ("Accepted for publication", "Accepted for publication"),
        ("In review", "In review"),
        ("Poster:", "Poster"),
    ):
        idx = rest.find(marker)
        if idx > 0:
            return clean(rest[:idx]), clean(rest[idx:])

    if status == "preprint" and rest.endswith("In review"):
        return clean(rest[: -len("In review")]), "In review"

    match = re.match(r"(?P<title>.+?\.)\s+(?P<journal>.+)", rest)
    if match:
        return clean(match.group("title")), clean(match.group("journal"))
    return clean(rest), ""


def parse_status_phrase_entry(block: str, marker: str, journal_label: str) -> tuple[str, str, str]:
    prefix, suffix = block.split(marker, 1)
    prefix = clean(prefix)
    suffix = clean(f"{marker}{suffix}")
    if ". " in prefix:
        authors, title = prefix.rsplit(". ", 1)
        return clean(authors), clean(title), suffix if journal_label else ""
    return "", prefix, suffix if journal_label else ""


def parse_authors_title_journal(block: str, status: str) -> tuple[str, str, str]:
    block = clean(block)

    if "Accepted for publication" in block:
        return parse_status_phrase_entry(block, "Accepted for publication", "Accepted for publication")
    if "In review" in block:
        return parse_status_phrase_entry(block, "In review", "In review")
    if "Proceedings of" in block:
        authors, rest = block.split("Proceedings of", 1)
        title, journal = parse_remainder(f"Proceedings of{rest}", status)
        return clean(authors), title, journal

    year_match = re.search(r"\((19|20)\d{2}\)\.?\s*", block)
    if year_match:
        authors = clean(block[:year_match.start()])
        title, journal = parse_remainder(block[year_match.end():], status)
        return authors, title, journal

    authors, rest = split_authors_and_rest(block)
    if authors:
        title, journal = parse_remainder(rest, status)
        return authors, title, journal

    title, journal = parse_remainder(block, status)
    return "", title, journal


def parse_section(lines: list[str], status: str, pub_type: str) -> list[dict]:
    entries = []
    for block in collect_entry_blocks(lines):
        doi = parse_doi(block)
        year = parse_year(block)
        pmid_match = PMID_RE.search(block)
        authors, title, journal = parse_authors_title_journal(block, status)

        if not title:
            continue

        entry = {
            "id": make_id(title),
            "title": title,
            "authors": clean(authors),
            "year": year,
            "journal": clean(journal),
            "doi": doi,
            "pmid": pmid_match.group(1) if pmid_match else "",
            "openalex_id": "",
            "status": status,
            "pub_type": pub_type,
            "source": "cv",
            "notes": "",
        }
        entries.append(entry)
    return entries


def parse_cv(cv_path: Path) -> list[dict]:
    text = get_text(cv_path)
    sections = split_into_sections(text)

    all_entries = []
    seen_ids = set()
    for section_name, (status, pub_type) in SECTION_CONFIG.items():
        for entry in parse_section(sections.get(section_name, []), status, pub_type):
            if entry["id"] in seen_ids:
                continue
            all_entries.append(entry)
            seen_ids.add(entry["id"])
    return all_entries


def write_csv(entries: list[dict], out_path: Path):
    with open(out_path, "w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=CSV_FIELDS)
        writer.writeheader()
        writer.writerows(entries)
    print(f"[✓] Wrote {len(entries)} entries → {out_path}")


def main():
    ap = argparse.ArgumentParser(description="Parse CV → publications CSV")
    ap.add_argument("--cv", required=True, help="Path to CV (PDF, .docx, .doc, .txt)")
    ap.add_argument("--out", default="publications.csv", help="Output CSV path")
    args = ap.parse_args()

    cv_path = Path(args.cv)
    out_path = Path(args.out)

    if not cv_path.exists():
        sys.exit(f"[ERROR] CV file not found: {cv_path}")

    print(f"[→] Parsing CV: {cv_path}")
    entries = parse_cv(cv_path)
    print(f"[→] Found {len(entries)} publication entries")
    write_csv(entries, out_path)


if __name__ == "__main__":
    main()
